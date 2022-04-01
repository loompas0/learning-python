# First Project of real Python program
# Creator : Loompas0 
# publish to github
# Date of creation : 12th of march 2022
# Version 0.1

# import from library docx import the Document
from datetime import date
from turtle import width
from docx import Document 
from docx.shared import Cm
import pyttsx3

# create a fuction for speak
def speak(text):
    pyttsx3.speak (text)

document = Document ()

# create some text in the document / cv
# input information for the document
first_name: str = input("What is your first name : ")
last_name : str = input("What is your last name : ")
date_of_Birth: date = input("What is your date of birth : ")
mobile_phone: str = input("What is your mobile phone number : ")
email: str = input("What is your email address : ")
about_me: str = input("Tell me about yourself : ")

# using speak function create above
speak ("Bonjour " + first_name + "Comment allez vous aujourd'hui ? ")
# Add a picture to our document
document.add_picture (
    "CV.jpg",
    width=Cm(3),
    height=Cm(4)
    )

# create a paragraph 1 in the document
document.add_paragraph (
    first_name + ' ' + last_name + " | " +
    str(date_of_Birth) + " | " +
    mobile_phone + " | " +
    email 
    )

# create some heading global information
document.add_heading ("About me")

# create a paragraph 2 in the document
document.add_paragraph (
    about_me
)

# create some heading Work Experience
document.add_heading ("Work Experience")

# create a paragraph  in the document using variable for experience
# loop for entering all experience until experience is == ""
while True:
    #       create a function 
    para = document.add_paragraph() 
    #       enter the variables for experiences
    # "commented" work_exp: str = input("What is your work experience : ")

    company: str = input ("What company : ")
    if company != "":
        from_date: str = input("Whait is your start date : ")
        end_date: str = input("Whait is your end date : ")
        experienceDetails: str = input (
        "Please could you describe your expererience at " 
        + company
        + " : ")
        #       print the experience step by step
        para.add_run(company + " ").bold = True # formatting bold
        para.add_run(
        from_date + " - " 
        + end_date + "\n").italic = True # formatting itallic
        para.add_run(experienceDetails)
    else:
        break
# End of the loop

# create some heading Skills
document.add_heading ("Skills")
# create a paragraph  in the document using variable for skills
# loop for entering all experience until skills is == ""
#       create a function 

while True:
#   enter the variables for skills
    skills: str = input ("What is your skills : ")
    if skills != "":
        para = document.add_paragraph() 
        para.style = "List Bullet" # formatting bullet

        #       print the skills bullet point
        para.add_run(skills) 
    else:
        break
# Footer of the document 
section = document.sections[0]
footer = section.footer
para = footer.paragraphs[0]
para.text = "CV automatically generated by Python Program"

# give a name and save the document 
nameOfFile: str = input ("Name of file without extension : ")
document.save(nameOfFile+".docx") 
